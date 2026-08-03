"""Build a template-derived, non-final paper development draft.

It intentionally contains only front-matter and data/method framing.  Final
Q1--Q4 result chapters are held back until Q3's human gate and Secondary test
complete; this prevents a document copy from masquerading as a frozen paper.
"""
from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
REFERENCE = ROOT / "论文.docx"
OUTPUT = ROOT / "paper" / "论文开发稿_前置章节_待冻结.docx"
FIGURE = ROOT / "outputs" / "development_figures" / "evidence_workflow_development.png"
MANIFEST = ROOT / "paper" / "开发稿构建清单.json"
REFERENCE_SHA256 = "f510ffe58f1c2c8a3cf4a9481038b777770daaf302c6b4f691a367b0858b6256"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_text(paragraph, text: str) -> None:
    """Replace text while preserving the first run's template formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_text(extra, "")


def main() -> None:
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("论文.docx 的 SHA-256 已变化；需重新蒸馏模板后再生成开发稿。")
    if not FIGURE.exists():
        raise FileNotFoundError(f"缺少开发流程图：{FIGURE}")

    document = Document(REFERENCE)
    paragraphs = document.paragraphs

    # Cover / abstract.
    set_text(paragraphs[0], "不同 SOC 区间快充策略与电池寿命研究")
    paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text(paragraphs[2], "开发稿（前置章节，待 Q3/Q4 结果冻结）。本文基于题目提供的锂离子电池循环老化数据，建立从数据深层审计、策略—寿命关联、早期运行校正到受可信域约束策略比较的证据流程。本版仅写入已验证的数据口径、问题分解、模型边界和可复现流程，不将尚未完成的 Q3 人工稳定性裁决、Secondary 压力测试或新策略 pilot 写成最终结论。")
    set_text(paragraphs[3], "问题一：按物理条码重建电芯级数据、审计局部异常，并给出寿命分布与重复策略的描述性一致性证据。")
    set_text(paragraphs[4], "问题二：在策略组折外验证下建立主效应与交互敏感性基线，严格区分条件统计关联、可信域与因果结论。")
    set_text(paragraphs[5], "问题三：以截止循环前的早期运行信息预测寿命与未来 SOH；k=5 用于筛查候选，k=100 用于开发集内正式校正候选，最终外部确认保留给 Secondary。")
    set_text(paragraphs[6], "问题四：将已有策略的开发池比较与新策略的试验确认分开；没有真实早期运行数据的新策略只能标记为 provisional。")
    set_text(paragraphs[8], "关键词：锂离子电池；分段快充；岭回归；早期健康预测；分组交叉验证；Pareto")

    # Problem restatement and analysis.
    set_text(paragraphs[13], "随着新能源汽车和储能系统快速发展，快充策略需要同时面对补能时间与循环寿命的权衡。题目数据记录了同一放电制度下不同两阶段恒流快充策略的循环老化过程。研究的重点不是把一个点预测包装成最终最优，而是在可追溯的数据口径下，区分策略层的条件关联、单电芯早期校正和仍待试验确认的新策略。")
    set_text(paragraphs[14], "原始数据由三个 MAT 文件及文献 Table 9 的官方寿命/策略记录组成。经审计后以 124 个物理条码为主体，使用官方 `cycle_life` 作为寿命标签；充电策略由 C1、SOC 切换比例 q 和 C2 表示。IR 的单位为 Ω，`chargetime` 的单位为 min；3.6 V 上限来自题目参考文献中的实验协议，不是题面额外硬约束。")
    set_text(paragraphs[16], "本文按“数据审计—策略关联—早期校正—受约束比较”的顺序回答四个问题，并用策略分组验证、字段—循环掩码和外部测试留置来控制信息泄漏。")
    set_text(paragraphs[17], "问题一：构建可追溯的电芯级数据集，描述寿命分布、策略覆盖与重复策略的跨分区一致性；不作因果归因。")
    set_text(paragraphs[18], "问题二：量化不同 SOC 阶段倍率与寿命的条件统计关联，比较主效应 Ridge 与二阶交互 Ridge，并给出可信域内的策略层预测代理。")
    set_text(paragraphs[19], "问题三：从容量、内阻、温度、充电时间及经审计的原始充电曲线中提取早期特征，预测未来 SOH 与寿命，比较观测窗口的时效—精度权衡。")
    set_text(paragraphs[20], "问题四：在已有策略与新策略两个信息条件下，分别进行开发池比较与试验设计；任何新策略在取得真实早期运行数据前均不输出正式 Pareto 推荐。")

    set_text(paragraphs[24], "问题一首先解决“什么数据可以被使用”。以 `barcode` 为物理主键拼接续测片段，使用 Table 9 官方寿命标签对齐策略，并保留所有电芯。原始曲线的异常不通过静默删整枚电芯处理，而是逐字段、逐循环形成掩码；随后用分位数与重抽样刻画寿命分布及重复策略的一致性。")
    set_text(paragraphs[25], "问题一的输出为后续模型提供统一样本边界、单位、循环索引与异常标记。由于相同策略在不同电芯和分区中仍存在显著差异，策略统计只能作为总体描述和后续条件关联的起点，不能替代个体化预测。")
    set_text(paragraphs[28], "问题二以策略参数为输入、自然对数寿命为目标。先使用主效应 Ridge 作为低复杂度基线，再用二阶交互 Ridge 检验 SOC 分段倍率的探索性关联；所有标准化、正则化和模型比较均置于 Train 内的 `policy_table9` 分组验证中。")
    set_text(paragraphs[30], "问题三把可获得信息的时点写入模型：截止 k 时只允许使用 cycle≤k 的数据。寿命和未来 SOH 采用相同的外层分组验证；k=5 的角色是尽早筛查，k=100 的角色是更充分的运行后校正候选。曲线增强特征只从通过六字段审计的充电点提取。")
    set_text(paragraphs[32], "问题四不直接在无约束网格上宣布最优点。对已有策略，只做开发池内的交叉拟合比较；对新策略，先用 Q2 在 raw/SOC 双空间可信域内提名，再经过真实 k=5 筛查、真实 k=100 校正和最终外部压力测试后，才可能进入正式 Pareto 集。")

    # Assumptions -- each is an explicit scope condition, not a causal claim.
    assumptions = [
        "1. 以 Table 9 对齐后的官方循环寿命作为电芯级寿命标签；原始 MAT 只读。",
        "2. 局部异常以字段—循环掩码处理，未通过掩码的字段不能进入相应特征，但不静默删除整枚电芯。",
        "3. 在 Train 的策略组折外比较中，策略—寿命关系只解释为条件统计关联，不解释为充电倍率的单独因果效应。",
        "4. 截止窗口 k 的预测只能使用 cycle≤k 的信息；训练、调参与模板构造均在外层训练折内完成。",
        "5. 模型仅适用于已审计数据覆盖的策略区域及其双空间可信邻域；3.6 V 仅作为参考文献实验协议的信息，不额外创造题外安全阈值。",
    ]
    for index, text in enumerate(assumptions, start=33):
        set_text(paragraphs[index], text)

    # Symbol table (first table is the template's symbol table).
    rows = [
        ("符号", "说明", "单位"),
        ("C1", "第一阶段充电倍率", "C-rate"),
        ("q", "倍率切换 SOC 比例（Q1_percent/100）", "—"),
        ("C2", "第二阶段充电倍率", "C-rate"),
        ("L_i", "Table 9 官方循环寿命", "cycle"),
        ("z_i=ln(L_i)", "自然对数寿命建模目标", "—"),
        ("k", "早期观测截止循环", "cycle"),
        ("SOH_nom", "放电容量/1.1 Ah", "—"),
        ("IR", "内阻字段", "Ω"),
        ("t_chg", "实测充电时间（chargetime）", "min"),
    ]
    for row, values in zip(document.tables[0].rows, rows):
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

    # Replace template preprocessing material with verified P0 protocol.
    preprocess = {
        42: "数据审计与预处理",
        43: "数据预处理的目标是建立可复现、可追溯且不泄漏的电芯级视图。原始三个 MAT 文件保持只读；文献 Table 9 提供官方寿命与策略对齐信息。所有派生表、掩码和图表均由脚本另存。",
        44: "物理主键、续测片段与官方标签",
        45: "三个原始批次共包含 140 条记录。以条码为主键拼接续测片段后，得到 124 个物理电芯；5 个续测片段并入相同物理电芯，11 个未进入 Table 9 的本地记录保留审计证据但不进入正式寿命分析。",
        46: "寿命标签不由本地序号截断推断，而使用 Table 9 的官方 `cycle_life`。策略参数、批次和寿命通过条码一对一对齐，并保留每条合并与排除记录的证据。",
        47: "字段—循环深层掩码",
        48: "对 t、Qc、I、V、T、Qd 逐字段检查长度一致性、NaN/Inf、时间倒序和物理边界。任何异常只屏蔽受影响字段或循环，不因少数局部异常删除整枚电芯。",
        49: "用于原始曲线的低维特征只取同时通过六字段审计的循环，并仅在 I>0.1 A 的充电点提取电压统计量；不对原始曲线静默插补。",
        50: "数据分区与泄漏控制",
        51: "Train、Primary、Secondary 分别含 41、43、40 个物理电芯。Q2/Q3 的模型选择、超参数搜索、标准化和 bootstrap 只读取 Train；Primary 只按冻结规则做一次受限观察，Secondary 留作最终独立压力测试。",
        52: "所有预测比较按 `policy_table9` 分组，避免相同策略同时出现在训练和测试折中。每个外层折内独立完成缺失处理、标准化、正则化选择及 SOH 模板构造。",
        53: "寿命与 SOH 指标",
        54: "寿命预测在 z=ln(L) 尺度计算 RMSE_log 和 MAE_log，再以 exp 反变换为循环寿命。字段名中的 log 为历史兼容名称，均不表示 log10。",
        55: "未来 SOH 误差按电芯先计算再等权汇总，避免单个寿命很长的电芯因循环点更多而主导评价。",
        56: "早期窗口与特征范围",
        57: "窗口 k∈{5,10,20,50,100}。当 k=5 时，实际可用的是 cycle 2–5 的 4 个循环点；特征包括容量、内阻、温度、充电时间的汇总、斜率与变化量。",
        58: "RAW challenger 另行加入三个经审计的充电电压特征，不与主线的汇总特征混淆。它只服务于最早筛查候选比较，不能据此自动宣称寿命预测已稳健优于基线。",
        59: "开发状态与后续锁定",
        60: "当前已完成 P0 深层审计、Q1 描述性重抽样、Q2 策略代理比较、Q3 联合/曲线增强候选比较，以及 Q4 已有策略的开发池敏感性分析。以下流程图仅说明证据流与状态，不是最终推荐图。",
    }
    for index, text in preprocess.items():
        set_text(paragraphs[index], text)

    # Use existing slots for the status note; delete unsupported template chapters.
    set_text(paragraphs[94], "后续结果章节的解锁条件")
    set_text(paragraphs[95], "本开发稿不包含最终 Q1—Q4 结果分析、最终策略或外部泛化结论。Q3 仍需由建模者完成稳定性和结果裁决；随后才能冻结数值，并在规则完全冻结后一次性运行 Secondary 压力测试。")
    set_text(paragraphs[96], "在此之前，Q3 的 M3R-k=5 只能作为早期曲线增强筛查候选，M2-k=100 只能作为开发集内误差最低的正式校正候选；Q4 的 4 个已有策略非支配点也只表示开发池案例。")
    set_text(paragraphs[97], "可复现证据见 `robustness/Q3/q3_robustness_report.md`、`robustness/Q4/q4_robustness_report.md` 与 `paper/audits/cross_media_consistency_audit.md`。")

    # Remove every remaining unsupported model/result/template chapter.
    for paragraph in list(document.paragraphs[98:]):
        remove_paragraph(paragraph)
    for table in list(document.tables[1:]):
        table._element.getparent().remove(table._element)

    document.add_picture(str(FIGURE), width=Inches(6.0))
    caption = document.add_paragraph("图 1  开发版证据流程与验证边界（非最终推荐）")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "图表标题" in [style.name for style in document.styles]:
        caption.style = "图表标题"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("生成过程改变了模板原件。")
    MANIFEST.write_text(json.dumps({
        "status": "development_front_matter_only",
        "reference": str(REFERENCE),
        "reference_sha256": REFERENCE_SHA256,
        "output": str(OUTPUT),
        "figure": str(FIGURE),
        "final_result_sections_included": False,
        "render_status": "blocked_no_libreoffice",
    }, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()

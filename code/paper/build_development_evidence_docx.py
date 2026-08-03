"""Render the non-final development evidence manuscript into the supplied template.

This is intentionally a development document.  Its first page and footer-level
notice prevent it from being mistaken for the final, Secondary-validated paper.
"""

from __future__ import annotations

import hashlib
import json
import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "论文.docx"
SOURCE = ROOT / "paper" / "开发证据稿_全篇_非最终.md"
OUTPUT = ROOT / "paper" / "论文开发证据稿_非最终.docx"
MANIFEST = ROOT / "paper" / "开发证据稿构建清单.json"
FROZEN_SOURCE = ROOT / "paper" / "论文冻结结果版_待提交前复核.md"
FROZEN_OUTPUT = ROOT / "paper" / "论文冻结结果版_待提交前复核.docx"
FROZEN_MANIFEST = ROOT / "paper" / "论文冻结结果版_构建清单.json"
FIGURES = {
    "数据审计": ROOT / "outputs" / "development_figures" / "evidence_workflow_development.png",
    "Q1": ROOT / "results" / "Q1" / "experiments" / "round1" / "figures" / "q1_01_lifetime_distribution.png",
    "Q2": ROOT / "results" / "Q2" / "experiments" / "round1" / "figures" / "m1_m2_oof_comparison.png",
    "Q3": ROOT / "results" / "Q3" / "experiments" / "round3_raw_curve_challenger" / "figures" / "q3_raw_curve_challenger.png",
    "Q4": ROOT / "results" / "Q4" / "experiments" / "existing_policy_round2_m2k100" / "figures" / "q4_existing_policy_pareto.png",
    "Q4流量": ROOT / "results" / "Q4" / "experiments" / "train_dry_run_round1" / "figures" / "q4_train_only_candidate_flow.png",
    "Q1策略": ROOT / "results" / "Q1" / "experiments" / "round1" / "figures" / "q1_07_policy_lifetime_top_bottom.png",
    "Q2稳健性": ROOT / "robustness" / "Q2" / "figures" / "q2_policy_block_bootstrap.png",
    "Q3外部": ROOT / "results" / "Secondary_final_pressure_test" / "figures" / "secondary_final_observed_predicted.png",
    "Q4试验": ROOT / "results" / "Q4" / "experiments" / "pilot_design_round1" / "figures" / "q4_pilot_representatives.png",
}
DOCX_ASSET_DIR = ROOT / "tmp" / "docx_raster_assets"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def delete_body(document: Document) -> None:
    body = document._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def style_run(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def clean_markup(text: str) -> str:
    """Keep the DOCX readable when the traceable source uses Markdown/TeX inline marks."""
    replacements = {
        "**": "",
        "$": "",
        "\\\\ln": "ln",
        "\\\\Delta": "Δ",
        "\\\\ge": "≥",
        "\\\\%": "%",
        "C_1": "C₁",
        "C_2": "C₂",
        "L_i": "Lᵢ",
        "z_i": "zᵢ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def add_paragraph(document: Document, text: str, *, centered: bool = False, bold: bool = False, size: float = 10.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21) if not centered else Pt(0)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(paragraph.add_run(clean_markup(text)), size=size, bold=bold)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(clean_markup(text))
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True


def add_formula(document: Document, latex: str) -> None:
    """Render the three verified development formulas in readable Word text.

    The source Markdown deliberately keeps LaTeX for traceability; python-docx
    has no native equation conversion, so these equivalent plain-text formulas
    avoid exposing raw TeX syntax in the handoff DOCX.
    """
    if "lVert" in latex:
        text = "min_{β₀, β} Σᵢ∈T (zᵢ − β₀ − xᵢᵀβ)² + λ‖β‖₂²"
    elif "widehat{SOH}" in latex:
        text = "ŜOHᵢ(c) = aᵢ(k) + [0.8 − aᵢ(k)] · [G(min{c/L̂ᵢ,1}) − G(k/L̂ᵢ)] / [0.8 − G(k/L̂ᵢ)]"
    elif "tau_" in latex:
        text = "τ₀–₈₀ = 60 · [q/C₁ + (0.8 − q)/C₂ᵉᶠᶠ] min"
    else:
        text = clean_markup(latex)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    style_run(paragraph.add_run(text), size=10.5)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.fullmatch(r"\|[\s:|\-]+\|", line):
            continue
        rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
    return rows


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_row = OxmlElement("w:tblHeader")
    header_row.set(qn("w:val"), "true")
    header_properties.append(header_row)
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            paragraph = table.cell(r, c).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.1
            style_run(paragraph.add_run(clean_markup(value)), size=8.5, bold=(r == 0))


def normalize_figure_for_docx(path: Path) -> Path:
    """Embed an opaque RGB copy so LibreOffice does not fail on RGBA PNG export.

    Original Chinese figures remain untouched; this is a renderer-only asset.
    """
    DOCX_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    target = DOCX_ASSET_DIR / f"{path.stem}_rgb.png"
    with Image.open(path) as image:
        if image.mode == "RGBA":
            background = Image.new("RGB", image.size, "white")
            background.paste(image, mask=image.getchannel("A"))
        else:
            background = image.convert("RGB")
        background.save(target, format="PNG", optimize=True)
    return target


def add_figure(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    asset = normalize_figure_for_docx(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = paragraph.add_run().add_picture(str(asset), width=Pt(390))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split("（", 1)[0])
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(caption_p.add_run(caption), size=9)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--frozen", action="store_true", help="生成冻结结果写作版，不覆盖开发证据稿。")
    args = parser.parse_args()
    source = FROZEN_SOURCE if args.frozen else SOURCE
    output = FROZEN_OUTPUT if args.frozen else OUTPUT
    manifest_path = FROZEN_MANIFEST if args.frozen else MANIFEST
    if not TEMPLATE.exists() or not source.exists():
        raise FileNotFoundError("模板或论文源稿不存在。")
    document = Document(TEMPLATE)
    delete_body(document)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = "｜冻结结果写作版｜提交前仍需版式与引用复核" if args.frozen else "｜开发证据版（非最终）｜仅供论文撰写与交接使用"
    style_run(footer.add_run(footer_text), size=8.5)

    title = "不同 SOC 区间快充策略与电池寿命的建模研究" if args.frozen else "不同 SOC 区间快充策略与电池寿命的开发期建模证据稿"
    subtitle = "冻结结果写作版（已纳入一次性 Secondary 压力测试；不将候选策略写成最终推荐）" if args.frozen else "开发证据版（非最终提交；待 Q3 裁决、Secondary 压力测试与新策略 pilot）"
    note = "说明：模型角色、窗口、特征、指标与 bootstrap 设置已冻结；Secondary 只读取一次，未据此调参或重选。" if args.frozen else "说明：本稿只整理已审计、可复现的开发证据。Primary 仅作一次受限确认；所有 Q3/Q4 结论均不等同于独立外部泛化或最终最优推荐。"
    add_paragraph(document, title, centered=True, bold=True, size=18)
    add_paragraph(document, subtitle, centered=True, size=11)
    add_paragraph(document, note, centered=True, size=9.5)

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    figure_inserted: set[str] = set()
    while index < len(lines):
        line = lines[index].strip()
        index += 1
        if not line or line.startswith(">") or line.startswith("# "):
            continue
        if line.startswith("## "):
            heading = line[3:]
            add_heading(document, heading, 1)
            if not args.frozen and heading.startswith("3 数据审计"):
                add_figure(document, FIGURES["数据审计"], "图 1  开发期证据流程与验证边界（非最终推荐）")
                figure_inserted.add("数据审计")
            elif not args.frozen and heading.startswith("4 Q1"):
                add_figure(document, FIGURES["Q1"], "图 2  正式分析电芯的循环寿命分布（数据来源：Table 9 对齐结果）")
                figure_inserted.add("Q1")
            elif not args.frozen and heading.startswith("5 Q2"):
                add_figure(document, FIGURES["Q2"], "图 3  Q2 策略分组折外预测：M1 主线与 M2 敏感性模型")
                figure_inserted.add("Q2")
            elif not args.frozen and heading.startswith("6 Q3"):
                add_figure(document, FIGURES["Q3"], "图 4  Q3 原始电压曲线候选模型的严格仅训练集比较（非外部验证）")
                figure_inserted.add("Q3")
            elif not args.frozen and heading.startswith("7 Q4"):
                add_figure(document, FIGURES["Q4"], "图 5  Q4 已有策略的开发池比较（非独立外部验证，非最终推荐）")
                add_figure(document, FIGURES["Q4流量"], "图 6  Q4 仅训练集候选流量与支持域诊断（不等于最优策略）")
                figure_inserted.add("Q4")
            elif args.frozen and heading.startswith("3 Q1"):
                add_figure(document, FIGURES["Q1"], "图 1  正式分析电芯的循环寿命分布")
                add_figure(document, FIGURES["Q1策略"], "图 2  具有重复支撑的策略组寿命比较（观察性对照）")
            elif args.frozen and heading.startswith("4 Q2"):
                add_figure(document, FIGURES["Q2"], "图 3  主效应 Ridge 与二阶交互 Ridge 的分组折外预测比较")
                add_figure(document, FIGURES["Q2稳健性"], "图 4  策略组块 bootstrap 下 M2 相对 M1 的误差差异")
            elif args.frozen and heading.startswith("5 Q3"):
                add_figure(document, FIGURES["Q3外部"], "图 5  Secondary 上冻结模型的实测—预测对比（外部压力测试）")
            elif args.frozen and heading.startswith("6 Q4"):
                add_figure(document, FIGURES["Q4流量"], "图 6  Train-only 双空间支持域候选筛选流程")
                add_figure(document, FIGURES["Q4试验"], "图 7  三类候选策略的 k=100 pilot 排程（待验证）")
            continue
        if line == "$$":
            formula_lines = []
            while index < len(lines) and lines[index].strip() != "$$":
                formula_lines.append(lines[index].strip())
                index += 1
            if index < len(lines):
                index += 1
            add_formula(document, " ".join(formula_lines))
            continue
        if line.startswith("| "):
            table_lines = [line]
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(document, parse_table(table_lines))
            continue
        add_paragraph(document, line)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    check = Document(output)
    manifest = {
        "status": "frozen_results_writing_docx" if args.frozen else "development_evidence_docx_nonfinal",
        "template": str(TEMPLATE),
        "template_sha256": sha256(TEMPLATE),
        "source_markdown": str(source),
        "source_sha256": sha256(source),
        "output": str(output),
        "output_sha256": sha256(output),
        "paragraph_count": len(check.paragraphs),
        "table_count": len(check.tables),
        "figure_count": 7 if args.frozen else 6,
        "final_submission_claim": False,
        "render_status": "structurally_verified; visual_render_pending_no_libreoffice",
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False))


if __name__ == "__main__":
    main()
